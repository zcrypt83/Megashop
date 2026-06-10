from pathlib import Path
from textwrap import wrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
ASSET_DIR = OUT_DIR / "plan_assets"
DOCX_PATH = OUT_DIR / "PLAN_DE_ACCION_MEGASHOP.docx"

NAVY = "172026"
TEAL = "0F8B8D"
TEAL_DARK = "0B6466"
GOLD = "F2B84B"
MUTED = "64717D"
LIGHT = "F3F6F8"
BORDER = "D8E1E8"
WHITE = "FFFFFF"
RED = "C84646"
GREEN = "16856B"


def rgb(value):
    return RGBColor.from_string(value)


def font(size=30, bold=False):
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def rounded_box(draw, xy, fill, outline=BORDER, radius=18, width=3):
    draw.rounded_rectangle(xy, radius=radius, fill="#" + fill, outline="#" + outline, width=width)


def text_box(draw, xy, title, lines, fill=WHITE, accent=TEAL, title_size=28, body_size=21):
    x1, y1, x2, y2 = xy
    rounded_box(draw, xy, fill)
    draw.rectangle((x1, y1, x1 + 12, y2), fill="#" + accent)
    draw.text((x1 + 30, y1 + 22), title, fill="#" + NAVY, font=font(title_size, True))
    y = y1 + 68
    for line in lines:
        for part in wrap(line, width=max(20, int((x2 - x1) / 13))):
            draw.text((x1 + 30, y), part, fill="#" + MUTED, font=font(body_size))
            y += body_size + 8
        y += 4


def arrow(draw, start, end, color=TEAL_DARK, width=5):
    draw.line((start, end), fill="#" + color, width=width)
    x2, y2 = end
    x1, y1 = start
    angle_dx = 14 if x2 >= x1 else -14
    angle_dy = 10 if y2 >= y1 else -10
    if abs(x2 - x1) >= abs(y2 - y1):
        points = [(x2, y2), (x2 - angle_dx, y2 - 10), (x2 - angle_dx, y2 + 10)]
    else:
        points = [(x2, y2), (x2 - 10, y2 - angle_dy), (x2 + 10, y2 - angle_dy)]
    draw.polygon(points, fill="#" + color)


def diagram_title(draw, title, subtitle):
    draw.text((70, 45), title, fill="#" + NAVY, font=font(38, True))
    draw.text((70, 98), subtitle, fill="#" + MUTED, font=font(22))
    draw.line((70, 140, 1530, 140), fill="#" + BORDER, width=3)


def create_architecture_diagram(path):
    image = Image.new("RGB", (1600, 900), "#EEF3F6")
    draw = ImageDraw.Draw(image)
    diagram_title(draw, "Arquitectura general", "Flujo principal de la plataforma Megashop")
    boxes = [
        ((90, 240, 350, 480), "Frontend", ["React 19 + Vite", "React Router", "Bootstrap 5"], GOLD),
        ((470, 240, 780, 480), "API REST", ["Spring Boot 3.5", "Spring MVC", "Spring Security"], TEAL),
        ((900, 190, 1210, 390), "MongoDB", ["Persistencia principal", "Documentos BSON", "Índices y agregaciones"], GREEN),
        ((900, 470, 1210, 670), "Redis", ["Sesiones y refresh", "Cache aside", "Tracking y rankings"], RED),
        ((1320, 320, 1530, 540), "DevOps", ["Docker Compose", "Contenedores", "Variables de entorno"], NAVY),
    ]
    for xy, title, lines, accent in boxes:
        text_box(draw, xy, title, lines, accent=accent)
    arrow(draw, (350, 360), (470, 360))
    arrow(draw, (780, 320), (900, 290))
    arrow(draw, (780, 400), (900, 560))
    arrow(draw, (1210, 290), (1320, 390))
    arrow(draw, (1210, 560), (1320, 480))
    draw.text((90, 760), "Principio: MongoDB conserva la fuente de verdad; Redis almacena datos temporales o reconstruibles.",
              fill="#" + NAVY, font=font(24, True))
    image.save(path)


def create_use_case_diagram(path):
    image = Image.new("RGB", (1600, 1120), "#EEF3F6")
    draw = ImageDraw.Draw(image)
    diagram_title(draw, "Diagrama de casos de uso", "Actores y capacidades principales")

    actor_columns = [
        ("Cliente", 80, 150, 540, GOLD),
        ("Administrador", 570, 150, 1030, TEAL),
        ("Repartidor", 1060, 150, 1520, GREEN),
    ]
    for name, x1, y1, x2, accent in actor_columns:
        rounded_box(draw, (x1, y1 + 25, x2, y1 + 105), WHITE, accent, 16, 4)
        bbox = draw.textbbox((0, 0), name, font=font(25, True))
        draw.text(((x1 + x2 - (bbox[2] - bbox[0])) / 2, y1 + 49), name,
                  fill="#" + accent, font=font(25, True))
        draw.line(((x1 + x2) / 2, y1 + 105, (x1 + x2) / 2, y1 + 145),
                  fill="#" + accent, width=4)

    columns = [
        (80, 540, GOLD, [
            "Registrarse / iniciar sesión",
            "Buscar y filtrar productos",
            "Gestionar carrito y checkout",
            "Consultar y cancelar pedido",
        ]),
        (570, 1030, TEAL, [
            "Gestionar clientes y productos",
            "Gestionar pedidos y entregas",
            "Consultar dashboard y reportes",
            "Administrar repartidores",
        ]),
        (1060, 1520, GREEN, [
            "Consultar entregas asignadas",
            "Actualizar estado del pedido",
            "Confirmar entrega",
            "Consultar ruta de reparto",
        ]),
    ]
    for x1, x2, accent, labels in columns:
        for index, label in enumerate(labels):
            y = 320 + index * 125
            draw.ellipse((x1 + 15, y, x2 - 15, y + 82), fill="#FFFFFF", outline="#" + accent, width=4)
            bbox = draw.textbbox((0, 0), label, font=font(19, True))
            draw.text(((x1 + x2 - (bbox[2] - bbox[0])) / 2, y + 28), label,
                      fill="#" + NAVY, font=font(19, True))
        draw.line(((x1 + x2) / 2, 295, (x1 + x2) / 2, 320), fill="#" + accent, width=4)

    draw.rounded_rectangle((260, 865, 1340, 1025), radius=18, fill="#FFFFFF", outline="#" + BORDER, width=3)
    draw.text((295, 890), "Casos compartidos", fill="#" + TEAL_DARK, font=font(25, True))
    shared = [
        ("Seguir entrega en tiempo real", 320, 945, GOLD),
        ("Consultar estado del pedido", 670, 945, TEAL),
        ("Mantener historial trazable", 1015, 945, GREEN),
    ]
    for label, x, y, accent in shared:
        draw.rounded_rectangle((x, y, x + 285, y + 54), radius=27, fill="#F3F6F8", outline="#" + accent, width=3)
        bbox = draw.textbbox((0, 0), label, font=font(16, True))
        draw.text((x + (285 - (bbox[2] - bbox[0])) / 2, y + 17), label,
                  fill="#" + NAVY, font=font(16, True))
    image.save(path)


def create_patterns_diagram(path):
    image = Image.new("RGB", (1600, 960), "#EEF3F6")
    draw = ImageDraw.Draw(image)
    diagram_title(draw, "Patrones de diseño", "Responsabilidades y dependencias del backend")

    layers = [
        ((120, 210, 1480, 330), "Controller / MVC", "Expone HTTP, aplica autorización y valida DTOs", GOLD),
        ((120, 370, 1480, 490), "Service Layer", "Concentra reglas de negocio y coordina MongoDB + Redis", TEAL),
        ((120, 530, 780, 660), "Repository Pattern", "Spring Data abstrae consultas y persistencia", GREEN),
        ((820, 530, 1480, 660), "Cache Aside", "Consulta Redis, resuelve desde MongoDB y repuebla cache", RED),
        ((120, 710, 780, 840), "DTO Pattern", "Records de entrada desacoplan HTTP del dominio", NAVY),
        ((820, 710, 1480, 840), "Dependency Injection", "Spring administra servicios, filtros y repositorios", TEAL_DARK),
    ]
    for xy, title, description, accent in layers:
        x1, y1, x2, y2 = xy
        rounded_box(draw, xy, WHITE, accent, 18, 4)
        draw.text((x1 + 28, y1 + 22), title, fill="#" + accent, font=font(27, True))
        draw.text((x1 + 28, y1 + 68), description, fill="#" + MUTED, font=font(21))
    arrow(draw, (800, 330), (800, 370))
    arrow(draw, (650, 490), (650, 530))
    arrow(draw, (950, 490), (950, 530))
    arrow(draw, (450, 660), (450, 710))
    arrow(draw, (1150, 660), (1150, 710))
    image.save(path)


def create_nosql_diagram(path):
    image = Image.new("RGB", (1600, 1260), "#EEF3F6")
    draw = ImageDraw.Draw(image)
    diagram_title(draw, "Modelo de datos no relacional", "Colecciones MongoDB, documentos embebidos y estructuras Redis")

    mongo_boxes = [
        ((80, 210, 430, 410), "usuarios", ["_id", "email (único)", "passwordHash", "rol"], TEAL),
        ((80, 470, 430, 700), "clientes", ["_id", "nombre, email", "direcciones [embebido]", "teléfono"], GREEN),
        ((610, 210, 990, 440), "productos", ["_id", "categoríaId", "precio, stock", "tags", "índice texto"], GOLD),
        ((610, 510, 990, 800), "pedidos", ["_id", "clienteId -> clientes", "repartidorId", "items [embebido]", "pago [embebido]", "estado, total"], TEAL),
        ((1170, 210, 1520, 410), "repartidores", ["_id", "nombre, teléfono", "placa, zona", "estado"], NAVY),
        ((1170, 500, 1520, 720), "historial_estados", ["_id", "pedidoId -> pedidos", "estado", "fecha, comentario"], RED),
    ]
    for xy, title, lines, accent in mongo_boxes:
        text_box(draw, xy, title, lines, accent=accent, title_size=25, body_size=19)

    arrow(draw, (430, 580), (610, 620), MUTED, 3)
    arrow(draw, (990, 620), (1170, 610), MUTED, 3)
    arrow(draw, (1170, 330), (990, 560), MUTED, 3)
    arrow(draw, (800, 440), (800, 510), MUTED, 3)

    draw.text((80, 860), "Redis - datos temporales y de alta frecuencia", fill="#" + NAVY, font=font(30, True))
    redis_items = [
        ("STRING", "session:{userId}, otp:{id}"),
        ("HASH", "order_status:{pedidoId}"),
        ("LIST", "client_orders:{clienteId}"),
        ("SET", "active_sessions"),
        ("SORTED SET", "most_viewed_products"),
        ("CACHE", "products:* y dashboard:resumen"),
    ]
    x, y = 80, 920
    for index, (kind, value) in enumerate(redis_items):
        col = index % 3
        row = index // 3
        x1 = x + col * 505
        y1 = y + row * 125
        rounded_box(draw, (x1, y1, x1 + 460, y1 + 92), WHITE, TEAL, 14, 3)
        draw.text((x1 + 20, y1 + 16), kind, fill="#" + TEAL_DARK, font=font(20, True))
        draw.text((x1 + 20, y1 + 51), value, fill="#" + MUTED, font=font(18))

    draw.text((80, 1175), "Regla: Redis nunca reemplaza la persistencia de pedidos, clientes o pagos.",
              fill="#" + NAVY, font=font(22, True))
    image.save(path)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_pr = table._tbl.tblPr
    tbl_w = table_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        table_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = table_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        table_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run(run, size=11, bold=False, color=NAVY, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = rgb(color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    set_run(run, 9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(NAVY)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in [
        ("Title", 28, NAVY, 0, 8),
        ("Subtitle", 14, MUTED, 0, 18),
        ("Heading 1", 16, TEAL_DARK, 16, 8),
        ("Heading 2", 13, TEAL_DARK, 12, 6),
        ("Heading 3", 12, NAVY, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = style_name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.text = "MEGASHOP  |  Plan de acción"
        set_run(header.runs[0], 9, True, MUTED)
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_page_number(section.footer.paragraphs[0])


def add_paragraph(doc, text, bold_lead=None, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run(lead, 11, True, NAVY)
        rest = p.add_run(text[len(bold_lead):])
        set_run(rest, 11, False, NAVY)
    else:
        run = p.add_run(text)
        set_run(run)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.45)
        p.paragraph_format.first_line_indent = Inches(-0.2)
        p.paragraph_format.space_after = Pt(5)
        set_run(p.add_run(item))


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.45)
        p.paragraph_format.first_line_indent = Inches(-0.2)
        p.paragraph_format.space_after = Pt(5)
        set_run(p.add_run(item))


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_widths(table, widths)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, TEAL_DARK)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.add_run(header), 9.5, True, WHITE)
    for row_index, row_data in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(row_data):
            if row_index % 2 == 1:
                set_cell_shading(cells[index], LIGHT)
            p = cells[index].paragraphs[0]
            set_run(p.add_run(str(value)), 9.3, False, NAVY)
    set_table_widths(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_figure(doc, path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Inches(6.45))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(3)
    cap.paragraph_format.space_after = Pt(8)
    run = cap.add_run(caption)
    set_run(run, 9, True, MUTED, True)


def page_break(doc):
    doc.add_page_break()


def build_document():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    diagrams = {
        "architecture": ASSET_DIR / "arquitectura.png",
        "use_cases": ASSET_DIR / "casos_uso.png",
        "patterns": ASSET_DIR / "patrones.png",
        "nosql": ASSET_DIR / "modelo_nosql.png",
    }
    create_architecture_diagram(diagrams["architecture"])
    create_use_case_diagram(diagrams["use_cases"])
    create_patterns_diagram(diagrams["patterns"])
    create_nosql_diagram(diagrams["nosql"])

    doc = Document()
    configure_document(doc)

    # Cover
    doc.add_paragraph().paragraph_format.space_after = Pt(80)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(kicker.add_run("PROYECTO DE BASES DE DATOS NO RELACIONALES"), 11, True, TEAL_DARK)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("PLAN DE ACCIÓN\nMEGASHOP")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Plataforma de comercio electrónico y delivery")

    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run("MongoDB  |  Redis  |  Spring Boot  |  React")
    set_run(run, 12, True, GOLD)
    doc.add_paragraph().paragraph_format.space_after = Pt(105)

    metadata = [
        ("Tipo de documento", "Plan de acción y diseño de solución"),
        ("Arquitectura", "Monolito modular preparado para microservicios"),
        ("Versión", "1.0"),
        ("Fecha", "9 de junio de 2026"),
    ]
    add_table(doc, ["Dato", "Detalle"], metadata, [2350, 7010])
    page_break(doc)

    # Contents
    doc.add_heading("Contenido", level=1)
    contents = [
        "1. Introducción",
        "2. Planteamiento del problema",
        "3. Objetivos",
        "4. Solución propuesta",
        "5. Tecnologías",
        "6. Diagramas de la solución",
        "7. Plan de acción",
        "8. Riesgos y mitigaciones",
        "9. Entregables y criterios de aceptación",
        "10. Conclusiones",
    ]
    add_numbered(doc, [item.split(". ", 1)[1] for item in contents])

    doc.add_heading("Resumen ejecutivo", level=1)
    add_paragraph(doc, "Megashop migrará la gestión de clientes, productos, pedidos, pagos y entregas desde hojas de cálculo hacia una plataforma web centralizada. La solución propuesta utiliza MongoDB para persistencia documental, Redis para información temporal y de alta frecuencia, Spring Boot para la API REST y React para el frontend.")
    add_paragraph(doc, "El plan organiza el trabajo en ocho fases: análisis, diseño, datos, backend, frontend, integración, pruebas y despliegue. Cada fase cuenta con actividades, responsables, entregables y criterios de cierre.")
    page_break(doc)

    doc.add_heading("1. Introducción", level=1)
    add_paragraph(doc, "Megashop es una iniciativa de modernización orientada a mejorar la venta y entrega de productos mediante pedidos en línea. Actualmente, la operación depende de hojas de cálculo separadas, lo que limita la trazabilidad, el crecimiento y la capacidad de responder en tiempo real.")
    add_paragraph(doc, "Este documento define la ruta de implementación del sistema, sus decisiones arquitectónicas, el modelo de datos no relacional y los controles necesarios para entregar una solución funcional, segura y mantenible.")

    doc.add_heading("2. Planteamiento del problema", level=1)
    add_paragraph(doc, "Situación actual:", bold_lead="Situación actual:")
    add_bullets(doc, [
        "La información de clientes, productos, pedidos, pagos y repartidores se mantiene en archivos separados.",
        "Existen registros duplicados y valores inconsistentes por falta de validaciones centralizadas.",
        "No hay un historial confiable de cambios de estado de los pedidos.",
        "Las consultas y reportes requieren consolidación manual.",
        "El seguimiento de entregas no opera en tiempo real.",
        "El modelo actual no escala de forma segura ante un mayor volumen de pedidos.",
    ])
    add_paragraph(doc, "Impacto:", bold_lead="Impacto:")
    add_bullets(doc, [
        "Mayor tiempo de atención y riesgo de errores operativos.",
        "Dificultad para medir ventas, inventario y desempeño de reparto.",
        "Experiencia inconsistente para clientes y administradores.",
        "Dependencia de procesos manuales sin controles de acceso adecuados.",
    ])

    doc.add_heading("3. Objetivos", level=1)
    doc.add_heading("3.1 Objetivo general", level=2)
    add_paragraph(doc, "Diseñar e implementar una plataforma empresarial de comercio electrónico y delivery que centralice la operación de Megashop mediante MongoDB, Redis, Spring Boot y React.")
    doc.add_heading("3.2 Objetivos específicos", level=2)
    add_bullets(doc, [
        "Centralizar la información y reducir duplicidades.",
        "Exponer una API REST versionada y protegida por roles.",
        "Implementar catálogo, carrito, pedidos, tracking y dashboard.",
        "Aplicar índices, documentos embebidos y agregaciones en MongoDB.",
        "Usar Redis para sesiones, cache, OTP, rankings y estados calientes.",
        "Automatizar el despliegue mediante Docker Compose.",
        "Generar documentación, dataset y pruebas reproducibles.",
    ])

    doc.add_heading("4. Solución propuesta", level=1)
    add_paragraph(doc, "La solución será un monolito modular preparado para separar dominios en microservicios cuando el crecimiento lo justifique. El frontend consumirá una API REST bajo `/api/v1`. MongoDB será la fuente de verdad y Redis almacenará información temporal o reconstruible.")
    add_figure(doc, diagrams["architecture"], "Figura 1. Arquitectura general propuesta.")

    doc.add_heading("5. Tecnologías", level=1)
    tech_rows = [
        ("Frontend", "React 19, Vite, React Router, Bootstrap 5, Lucide", "Interfaz responsive y componentes reutilizables"),
        ("Backend", "Java 21, Spring Boot 3.5, Spring MVC", "API REST y reglas de negocio"),
        ("Seguridad", "Spring Security, JWT, BCrypt, Bean Validation", "Autenticación, roles y validación"),
        ("Persistencia", "MongoDB 7, Spring Data MongoDB", "Documentos BSON, índices y agregaciones"),
        ("Cache", "Redis 7, Spring Data Redis", "Sesiones, tracking, cache y rankings"),
        ("DevOps", "Docker Compose, Maven, Eclipse Temurin", "Build y despliegue reproducible"),
        ("Pruebas", "JUnit 5, Spring Boot Test, Postman", "Verificación técnica y funcional"),
    ]
    add_table(doc, ["Capa", "Tecnologías", "Uso"], tech_rows, [1500, 3900, 3960])
    page_break(doc)

    doc.add_heading("6. Diagramas de la solución", level=1)
    doc.add_heading("6.1 Casos de uso", level=2)
    add_paragraph(doc, "Los actores principales son Cliente, Administrador y Repartidor. Las acciones se limitan mediante RBAC y cada cambio de pedido conserva trazabilidad.")
    add_figure(doc, diagrams["use_cases"], "Figura 2. Diagrama de casos de uso.")

    doc.add_heading("6.2 Patrones de diseño", level=2)
    add_paragraph(doc, "El backend aplica MVC, Service Layer, Repository, DTO, Dependency Injection y Cache Aside. Esta combinación reduce acoplamiento y facilita pruebas, mantenimiento y extracción futura de módulos.")
    add_figure(doc, diagrams["patterns"], "Figura 3. Patrones y capas del backend.")

    doc.add_heading("6.3 Modelo de base de datos no relacional", level=2)
    add_paragraph(doc, "Los pedidos embeben items y datos de pago para optimizar la lectura. Las relaciones con clientes y repartidores se representan mediante identificadores. El historial se mantiene en una colección independiente para auditoría.")
    add_figure(doc, diagrams["nosql"], "Figura 4. Modelo MongoDB y estructuras Redis.")

    doc.add_heading("7. Plan de acción", level=1)
    phase_rows = [
        ("1", "Análisis", "Revisar alcance, actores, reglas, datos actuales y riesgos.", "Arquitecto / Analista", "Semana 1", "Backlog y alcance aprobados"),
        ("2", "Diseño", "Definir arquitectura, API, seguridad, UX y modelo NoSQL.", "Arquitecto / UX", "Semana 2", "Diagramas y contratos"),
        ("3", "Datos", "Crear colecciones, índices, scripts y dataset.", "DBA MongoDB / Redis", "Semana 3", "Base reproducible"),
        ("4", "Backend", "Implementar auth, clientes, productos, pedidos y dashboard.", "Backend Java", "Semanas 4-5", "API compilada y probada"),
        ("5", "Frontend", "Construir catálogo, checkout, tracking, login y admin.", "Frontend React", "Semanas 5-6", "SPA responsive"),
        ("6", "Integración", "Conectar API, cache, sesiones, errores y estados.", "Equipo full stack", "Semana 7", "Flujos integrados"),
        ("7", "Calidad", "Ejecutar pruebas, seguridad, rendimiento y accesibilidad.", "QA / Seguridad", "Semana 8", "Informe de pruebas"),
        ("8", "Despliegue", "Contenerizar, cargar datos, documentar y presentar.", "DevOps / Equipo", "Semana 9", "Release demostrable"),
    ]
    add_table(doc, ["Fase", "Nombre", "Actividades", "Responsable", "Plazo", "Resultado"], phase_rows,
              [600, 1150, 3450, 1450, 1050, 1660])

    doc.add_heading("7.1 Priorización", level=2)
    add_bullets(doc, [
        "Prioridad alta: autenticación, productos, pedidos, estados e índices.",
        "Prioridad media: dashboard, favoritos, notificaciones y reportes.",
        "Prioridad evolutiva: OTP, promociones, mapas reales y observabilidad avanzada.",
    ])

    doc.add_heading("7.2 Hitos de control", level=2)
    milestone_rows = [
        ("H1", "Diseño aprobado", "Arquitectura, API y modelo de datos revisados"),
        ("H2", "MVP backend", "Auth, productos y pedidos operativos"),
        ("H3", "MVP frontend", "Compra y seguimiento navegables"),
        ("H4", "Integración completa", "MongoDB y Redis conectados"),
        ("H5", "Entrega", "Docker, pruebas, documentación y demostración"),
    ]
    add_table(doc, ["Hito", "Nombre", "Criterio de salida"], milestone_rows, [900, 2500, 5960])

    doc.add_heading("8. Riesgos y mitigaciones", level=1)
    risk_rows = [
        ("Credenciales o secretos expuestos", "Alta", "Variables de entorno, rotación y gestor de secretos"),
        ("Datos duplicados en la migración", "Alta", "Índices únicos, limpieza y validación previa"),
        ("Redis no disponible", "Media", "Degradar a MongoDB y reconstruir cache"),
        ("Consultas lentas", "Media", "Índices, explain, límites y cache aside"),
        ("Cambios de alcance", "Media", "Backlog priorizado y control por hitos"),
        ("Fallas responsive o accesibilidad", "Media", "QA en tres viewports y WCAG AA"),
    ]
    add_table(doc, ["Riesgo", "Nivel", "Mitigación"], risk_rows, [3000, 1100, 5260])

    doc.add_heading("9. Entregables y criterios de aceptación", level=1)
    add_bullets(doc, [
        "Backend Spring Boot compilable y API disponible bajo `/api/v1`.",
        "Frontend React responsive con catálogo, checkout, seguimiento, login y dashboard.",
        "MongoDB con colecciones, índices, dataset y consultas avanzadas.",
        "Redis con sesiones, tracking, cache y ranking.",
        "Docker Compose con cuatro servicios.",
        "Colección Postman, documentación técnica y prototipo Figma Make.",
    ])

    acceptance_rows = [
        ("Funcionalidad", "Los flujos principales se completan sin errores bloqueantes."),
        ("Seguridad", "Las rutas privadas exigen JWT y respetan roles."),
        ("Datos", "No existen correos duplicados y los estados conservan historial."),
        ("Rendimiento", "Catálogo y dashboard usan cache; consultas usan índices."),
        ("UX", "No hay solapamientos; navegación usable en desktop, tablet y móvil."),
        ("Despliegue", "`docker compose up --build` levanta la solución."),
    ]
    add_table(doc, ["Dimensión", "Criterio"], acceptance_rows, [2200, 7160])

    doc.add_heading("10. Conclusiones", level=1)
    add_paragraph(doc, "El plan transforma una operación fragmentada en una plataforma centralizada y trazable. La combinación de MongoDB y Redis responde a necesidades diferentes: persistencia flexible y procesamiento rápido de información temporal.")
    add_paragraph(doc, "La arquitectura por capas, la seguridad basada en JWT y el despliegue con contenedores reducen riesgos técnicos y dejan una base preparada para crecimiento, integración con servicios externos y futura separación por dominios.")

    doc.add_heading("Anexo A. Prototipo de frontend", level=1)
    add_paragraph(doc, "El archivo `prototype/megashop-frontend.make` contiene la especificación de pantallas, sistema visual, componentes, interacciones, estados responsive y criterios de aceptación para reproducir el prototipo en Figma Make.")

    doc.core_properties.title = "Plan de Acción Megashop"
    doc.core_properties.subject = "Arquitectura, diagramas y plan de implementación"
    doc.core_properties.author = "Proyecto Megashop"
    doc.core_properties.keywords = "Megashop, MongoDB, Redis, Spring Boot, React, plan de acción"
    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    result = build_document()
    print(result)
